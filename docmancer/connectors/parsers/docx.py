from __future__ import annotations

from pathlib import Path

from docmancer.core.models import Document


class DOCXLoader:
    supported_extensions = [".docx"]
    chunking_strategy = "paragraph"

    def load(self, path: Path) -> Document:
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise ImportError("DOCX loader unavailable; reinstall docmancer (python-docx ships in core).") from exc

        doc = WordDocument(str(path))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                level_text = style_name.replace("heading", "").strip()
                level = int(level_text) if level_text.isdigit() else 2
                lines.append(f"{'#' * max(1, min(level, 6))} {text}")
            else:
                lines.append(text)

        content = "\n\n".join(lines).strip()
        if not content:
            raise ValueError(f"No extractable text found in DOCX: {path}")
        core = doc.core_properties
        return Document(source=str(path), content=content, metadata={"format": "docx", "title": core.title or path.stem})
